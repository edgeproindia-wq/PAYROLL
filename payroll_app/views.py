from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from decimal import Decimal
from .models import Employee, SalaryStructure, Attendance, LeaveRequest, Reimbursement, EmailVerificationToken, EmailOTPVerification
from .forms import EmployeeForm, SalaryStructureForm, AttendanceForm, LeaveRequestForm, ReimbursementForm, PayrollRunForm
from .models import PayrollRun, PayrollRunLine, InvestmentDeclaration
import csv
from django.http import HttpResponse


def generic_page(request, template_path):
    return render(request, f'{template_path}.html')


def dashboard(request):
    from django.db.models import Count
    import json
    dept_data = Employee.objects.filter(employment_status='ACTIVE').values('department').annotate(c=Count('id')).order_by('-c')
    dept_labels = json.dumps([d['department'] or 'Unassigned' for d in dept_data])
    dept_values = json.dumps([d['c'] for d in dept_data])
    status_data = Employee.objects.values('employment_status').annotate(c=Count('id'))
    status_map = {'ACTIVE': 'Active', 'ON_LEAVE': 'On Leave', 'RESIGNED': 'Resigned', 'TERMINATED': 'Terminated'}
    status_labels = json.dumps([status_map.get(s['employment_status'], s['employment_status']) for s in status_data])
    status_values = json.dumps([s['c'] for s in status_data])
    run_status_data = PayrollRun.objects.values('status').annotate(c=Count('id'))
    run_status_map = {'DRAFT': 'Draft', 'VALIDATED': 'Validated', 'APPROVED': 'Approved', 'RELEASED': 'Released'}
    run_labels = json.dumps([run_status_map.get(r['status'], r['status']) for r in run_status_data])
    run_values = json.dumps([r['c'] for r in run_status_data])
    salary_generated_count = PayrollRun.objects.filter(lines__isnull=False).distinct().count()
    latest_payroll_run = PayrollRun.objects.first()
    latest_run_net_pay = 0
    latest_run_employee_count = 0
    if latest_payroll_run:
        run_lines = latest_payroll_run.lines.all()
        latest_run_net_pay = sum(l.net_pay for l in run_lines)
        latest_run_employee_count = run_lines.count()
    pending_investment_declarations = InvestmentDeclaration.objects.filter(is_verified=False).count()
    recent_leaves = LeaveRequest.objects.select_related('employee').order_by('-id')[:5]
    recent_reimbursements = Reimbursement.objects.select_related('employee').order_by('-id')[:5]

    from django.db.models import Sum
    cost_data = PayrollRunLine.objects.filter(payroll_run__status='RELEASED').values('payroll_run__month').annotate(total=Sum('net_pay')).order_by('payroll_run__id')
    cost_labels = json.dumps([c['payroll_run__month'] for c in cost_data])
    cost_values = json.dumps([float(c['total']) for c in cost_data])

    context = {
        'dept_labels': dept_labels,
        'dept_values': dept_values,
        'status_labels': status_labels,
        'status_values': status_values,
        'run_labels': run_labels,
        'run_values': run_values,
        'total_employees': Employee.objects.filter(employment_status='ACTIVE').count(),
        'total_salary_structures': SalaryStructure.objects.count(),
        'salary_generated_count': salary_generated_count,
        'pending_leave': LeaveRequest.objects.filter(status='PENDING').count(),
        'pending_reimbursements': Reimbursement.objects.filter(status='PENDING').count(),
        'latest_payroll_run': latest_payroll_run,
        'latest_run_net_pay': latest_run_net_pay,
        'latest_run_employee_count': latest_run_employee_count,
        'pending_investment_declarations': pending_investment_declarations,
        'pending_approvals': PayrollRun.objects.filter(status='VALIDATED').count(),
        'recent_leaves': recent_leaves,
        'recent_reimbursements': recent_reimbursements,
        'cost_labels': cost_labels,
        'cost_values': cost_values,
    }
    return render(request, 'Dashboard.html', context)


def dashboard_payroll_cost_export_excel(request):
    import openpyxl
    from openpyxl.utils import get_column_letter
    from django.db.models import Sum

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Payroll Cost Summary"
    headers = ['Month', 'Total Net Payout']
    ws.append(headers)

    cost_data = PayrollRunLine.objects.filter(payroll_run__status='RELEASED').values('payroll_run__month').annotate(total=Sum('net_pay')).order_by('payroll_run__id')
    for c in cost_data:
        ws.append([c['payroll_run__month'], float(c['total'])])

    for i, header in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(i)].width = max(18, len(header) + 4)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="Payroll_Cost_Summary.xlsx"'
    wb.save(response)
    return response


def employee_master_export_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="employees.csv"'
    writer = csv.writer(response)
    writer.writerow(['Code', 'Name', 'Email', 'Department', 'Designation', 'Status', 'Date of Joining'])
    for emp in Employee.objects.all():
        writer.writerow([emp.employee_code, emp.full_name, emp.email, emp.department, emp.designation, emp.get_employment_status_display(), emp.date_of_joining])
    return response


def employee_master_import_xlsx(request):
    import openpyxl
    from datetime import date

    STATUS_MAP = {'working': 'ACTIVE', 'active': 'ACTIVE', 'terminated': 'TERMINATED', 'abscond': 'TERMINATED', 'resigned': 'RESIGNED', 'on leave': 'ON_LEAVE'}

    HEADER_ALIASES = {
        'code': ['emp code', 'employee code', 'code', 'emp id', 'employee id'],
        'name': ['name', 'employee name', 'full name'],
        'department': ['dept', 'department', 'team'],
        'designation': ['designation', 'title', 'role'],
        'doj': ['date of joining', 'doj', 'joining date'],
        'status': ['status', 'employment status'],
    }

    def find_col(headers, aliases):
        for i, h in enumerate(headers):
            if h and str(h).strip().lower() in aliases:
                return i
        return None

    if request.method == 'POST' and request.FILES.get('xlsx_file'):
        f = request.FILES['xlsx_file']
        wb = openpyxl.load_workbook(f, data_only=True)
        ws = wb.active

        header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        col_map = {key: find_col(header_row, aliases) for key, aliases in HEADER_ALIASES.items()}

        missing_required = [k for k in ('code', 'name') if col_map[k] is None]
        if missing_required:
            found = ', '.join(str(h) for h in header_row if h)
            messages.error(request, f"Import failed: could not find required column(s) {missing_required} in the uploaded file. Columns found: {found}")
            return redirect('employee_master')

        rows = list(ws.iter_rows(min_row=2, values_only=True))
        created, updated, skipped = 0, 0, 0

        def get(row, key):
            idx = col_map[key]
            if idx is None or idx >= len(row):
                return None
            return row[idx]

        for row in rows:
            code = get(row, 'code')
            name = get(row, 'name')
            if not code or not name:
                skipped += 1
                continue

            name_parts = str(name).strip().split(' ', 1)
            first_name = name_parts[0]
            last_name = name_parts[1] if len(name_parts) > 1 else ''

            department = get(row, 'department') or 'General'
            designation = get(row, 'designation') or 'Not Specified'
            doj = get(row, 'doj')
            status = get(row, 'status')
            emp_status = STATUS_MAP.get(str(status).strip().lower(), 'ACTIVE') if status else 'ACTIVE'

            obj, was_created = Employee.objects.update_or_create(
                employee_code=str(code).strip(),
                defaults={
                    'first_name': first_name,
                    'last_name': last_name,
                    'department': str(department),
                    'designation': str(designation),
                    'employment_status': emp_status,
                    'date_of_joining': doj.date() if hasattr(doj, 'date') else (doj or date(2025, 1, 1)),
                    'email': f"{str(code).lower().replace(' ', '')}@edgepro.local",
                },
            )
            created += 1 if was_created else 0
            updated += 0 if was_created else 1

        messages.success(request, f"Import complete. Created={created} Updated={updated} Skipped={skipped}")
        return redirect('employee_master')
    messages.error(request, "No file uploaded.")
    return redirect('employee_master')


def employee_master(request):
    from django.core.paginator import Paginator
    max_number = 0
    for code in Employee.objects.values_list('employee_code', flat=True):
        digits = ''.join(filter(str.isdigit, code or ''))
        if digits:
            max_number = max(max_number, int(digits))
    next_number = max_number + 1
    number_options = [str(n).zfill(3) for n in range(next_number, next_number + 20)]
    if request.method == 'POST':
        post_data = request.POST.copy()
        prefix = post_data.get('code_prefix', 'EPRO')
        number = post_data.get('employee_code', '')
        post_data['employee_code'] = f'{prefix}{number}'
        form = EmployeeForm(post_data)
        if form.is_valid():
            form.save()
            return redirect('employee_master')
    else:
        form = EmployeeForm()
    all_employees = Employee.objects.numeric_order()
    existing_codes = Employee.objects.values_list('employee_code', flat=True).order_by('employee_code')
    paginator = Paginator(all_employees, 50)
    page_number = request.GET.get('page')
    employees = paginator.get_page(page_number)
    return render(request, 'Employee Master.html', {'employees': employees, 'form': form, 'existing_codes': existing_codes, 'number_options': number_options})


def employee_edit(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    if request.method == 'POST':
        form = EmployeeForm(request.POST, instance=employee)
        if form.is_valid():
            form.save()
            return redirect('employee_master')
    else:
        form = EmployeeForm(instance=employee)
    return render(request, 'Employee Master.html', {
        'employees': Employee.objects.all(), 'form': form, 'edit_employee': employee,
    })


def employee_delete(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    employee.delete()
    return redirect('employee_master')


def salary_structure(request):
    import json
    if request.method == 'POST':
        form = SalaryStructureForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('salary_structure')
    else:
        form = SalaryStructureForm()
    from django.core.paginator import Paginator
    all_structures = SalaryStructure.objects.select_related('employee').all()
    paginator = Paginator(all_structures, 10)
    page_number = request.GET.get('page')
    structures = paginator.get_page(page_number)
    all_structures_for_chart = SalaryStructure.objects.select_related('employee').all()[:15]
    chart_labels = json.dumps([s.employee.full_name for s in all_structures_for_chart])
    chart_basic = json.dumps([float(s.basic) for s in all_structures_for_chart])
    chart_hra = json.dumps([float(s.hra) for s in all_structures_for_chart])
    return render(request, 'Salary Structure.html', {'form': form, 'structures': structures, 'chart_labels': chart_labels, 'chart_basic': chart_basic, 'chart_hra': chart_hra})


def salary_components(request):
    if request.method == 'POST':
        form = SalaryComponentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('salary_components')
    else:
        form = SalaryComponentForm()
    components = SalaryComponent.objects.all()
    return render(request, 'Salary Components.html', {'form': form, 'components': components})


def salary_templates(request):
    from .forms import SalaryTemplateForm
    if request.method == 'POST':
        form = SalaryTemplateForm(request.POST)
        if form.is_valid():
            template = form.save()
            for comp in SalaryComponent.objects.filter(is_active=True):
                value = request.POST.get(f'component_{comp.id}')
                if value:
                    SalaryTemplateComponent.objects.create(template=template, component=comp, value=value)
            return redirect('salary_templates')
    else:
        form = SalaryTemplateForm()
    templates = SalaryTemplate.objects.prefetch_related('components__component').all()
    all_components = SalaryComponent.objects.filter(is_active=True)
    return render(request, 'Salary Templates.html', {'form': form, 'templates': templates, 'all_components': all_components})


def employee_salary_components(request):
    from django.core.paginator import Paginator
    if request.method == 'POST':
        form = EmployeeSalaryComponentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('employee_salary_components')
    else:
        form = EmployeeSalaryComponentForm()
    all_assignments = EmployeeSalaryComponent.objects.select_related('employee', 'component').filter(is_active=True).order_by('employee__employee_code')
    paginator = Paginator(all_assignments, 15)
    page_number = request.GET.get('page')
    assignments = paginator.get_page(page_number)
    return render(request, 'Employee Salary Components.html', {'form': form, 'assignments': assignments})


def attendance(request):
    import json
    if request.method == 'POST':
        form = AttendanceForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('attendance')
    else:
        form = AttendanceForm()
    from django.core.paginator import Paginator
    all_records = Attendance.objects.select_related('employee').all()
    paginator = Paginator(all_records, 10)
    page_number = request.GET.get('page')
    records = paginator.get_page(page_number)
    from django.db.models import Count
    status_counts = Attendance.objects.values('status').annotate(c=Count('id'))
    status_map = {'PRESENT': 'Present', 'ABSENT': 'Absent', 'HALF_DAY': 'Half Day', 'LEAVE': 'On Leave'}
    chart_labels = json.dumps([status_map.get(s['status'], s['status']) for s in status_counts])
    chart_values = json.dumps([s['c'] for s in status_counts])
    return render(request, 'Attendance.html', {'form': form, 'records': records, 'chart_labels': chart_labels, 'chart_values': chart_values})


def leave_management(request):
    import json
    if request.method == 'POST':
        form = LeaveRequestForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('leave_management')
    else:
        form = LeaveRequestForm()
    from django.core.paginator import Paginator
    all_leave = LeaveRequest.objects.select_related('employee').all()
    paginator = Paginator(all_leave, 10)
    page_number = request.GET.get('page')
    leave_requests = paginator.get_page(page_number)
    from django.db.models import Count
    type_counts = LeaveRequest.objects.values('leave_type').annotate(c=Count('id'))
    type_map = dict(LeaveRequest.LEAVE_TYPE_CHOICES)
    chart_labels = json.dumps([type_map.get(t['leave_type'], t['leave_type']) for t in type_counts])
    chart_values = json.dumps([t['c'] for t in type_counts])
    return render(request, 'Leave Management.html', {'form': form, 'leave_requests': leave_requests, 'chart_labels': chart_labels, 'chart_values': chart_values})


def reimbursement(request):
    if request.method == 'POST':
        form = ReimbursementForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('reimbursement')
    else:
        form = ReimbursementForm()
    from django.core.paginator import Paginator
    all_reimb = Reimbursement.objects.select_related('employee').all()
    paginator = Paginator(all_reimb, 10)
    page_number = request.GET.get('page')
    reimbursements = paginator.get_page(page_number)
    return render(request, 'Reimbursement.html', {'form': form, 'reimbursements': reimbursements})


def statutory_compliance(request):
    import json
    rows = []
    for ss in SalaryStructure.objects.select_related('employee').all():
        pf = ss.basic * Decimal('0.12')
        esi = ss.gross_salary * Decimal('0.0075') if ss.gross_salary <= 21000 else Decimal('0')
        rows.append({'employee': ss.employee, 'basic': ss.basic, 'gross': ss.gross_salary, 'pf': round(pf, 2), 'esi': round(esi, 2)})
    chart_labels = json.dumps([r['employee'].full_name for r in rows])
    chart_pf = json.dumps([float(r['pf']) for r in rows])
    chart_esi = json.dumps([float(r['esi']) for r in rows])
    return render(request, 'Tax and Compliance/Statutory Compliance.html', {'rows': rows, 'chart_labels': chart_labels, 'chart_pf': chart_pf, 'chart_esi': chart_esi})


def investment_declaration(request):
    from .models import InvestmentDeclaration
    from .forms import InvestmentDeclarationForm
    if request.method == 'POST':
        form = InvestmentDeclarationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('investment_declaration')
    else:
        form = InvestmentDeclarationForm()
    declarations = InvestmentDeclaration.objects.select_related('employee').all()
    return render(request, 'Income Tax Management/Investment declartion.html', {'form': form, 'declarations': declarations})


def income_tax(request):
    import json
    rows = []
    for ss in SalaryStructure.objects.select_related('employee').all():
        annual_gross = ss.gross_salary * Decimal('12')
        if annual_gross <= 300000:
            tds_annual = 0
        elif annual_gross <= 700000:
            tds_annual = (annual_gross - Decimal('300000')) * Decimal('0.05')
        else:
            tds_annual = Decimal('400000') * Decimal('0.05') + (annual_gross - Decimal('700000')) * Decimal('0.10')
        rows.append({'employee': ss.employee, 'annual_gross': round(annual_gross, 2), 'tds_annual': round(tds_annual, 2), 'tds_monthly': round(tds_annual / 12, 2)})
    chart_labels = json.dumps([r['employee'].full_name for r in rows])
    chart_values = json.dumps([float(r['tds_monthly']) for r in rows])
    return render(request, 'Tax and Compliance/Income Tax.html', {'rows': rows, 'chart_labels': chart_labels, 'chart_values': chart_values})


def compliance_reports(request):
    import json
    structures = SalaryStructure.objects.all()
    total_pf = sum(s.basic * Decimal('0.12') for s in structures)
    total_esi = sum(s.gross_salary * Decimal('0.0075') for s in structures if s.gross_salary <= 21000)
    total_gross = sum(s.gross_salary for s in structures)
    chart_labels = json.dumps(['PF', 'ESI', 'Gross'])
    chart_values = json.dumps([float(round(total_pf, 2)), float(round(total_esi, 2)), float(round(total_gross, 2))])
    return render(request, 'Tax and Compliance/Compliance Reports.html', {
        'total_pf': round(total_pf, 2), 'total_esi': round(total_esi, 2), 'total_gross': round(total_gross, 2), 'employee_count': structures.count(),
        'chart_labels': chart_labels, 'chart_values': chart_values,
    })


def total_employees_report(request):
    import json
    from django.db.models import Count, Q
    dept_summary_qs = Employee.objects.values('department').annotate(
        total=Count('id'),
        active=Count('id', filter=Q(employment_status='ACTIVE'))
    ).order_by('department')
    dept_summary = [{'department': row['department'] or '(No Department)', 'total': row['total'], 'active': row['active']} for row in dept_summary_qs]
    chart_labels = json.dumps([d['department'] for d in dept_summary])
    chart_values = json.dumps([d['total'] for d in dept_summary])
    context = {
        'total': Employee.objects.count(),
        'active': Employee.objects.filter(employment_status='ACTIVE').count(),
        'on_leave': Employee.objects.filter(employment_status='ON_LEAVE').count(),
        'resigned': Employee.objects.filter(employment_status='RESIGNED').count(),
        'dept_summary': dept_summary,
        'employees': Employee.objects.all(),
        'chart_labels': chart_labels,
        'chart_values': chart_values,
    }
    return render(request, 'Payroll/Total Employees.html', context)


def new_joiners_report(request):
    from datetime import date, timedelta
    cutoff = date.today() - timedelta(days=90)
    employees = Employee.objects.filter(date_of_joining__gte=cutoff).order_by('-date_of_joining')
    return render(request, 'Payroll/New Joiners.html', {'employees': employees, 'cutoff': cutoff})


def payroll_cost_report(request):
    import json
    from django.db.models import Sum, Count, F
    dept_costs_qs = SalaryStructure.objects.values('employee__department').annotate(
        employee_count=Count('id'),
        total_cost=Sum(F('basic') + F('hra') + F('conveyance') + F('special_allowance'))
    ).order_by('employee__department')
    dept_costs = [{'department': row['employee__department'], 'employee_count': row['employee_count'], 'total_cost': row['total_cost'] or 0} for row in dept_costs_qs]
    overall_total = sum(d['total_cost'] for d in dept_costs)
    chart_labels = json.dumps([d['department'] for d in dept_costs])
    chart_values = json.dumps([float(d['total_cost']) for d in dept_costs])
    return render(request, 'Payroll/Payroll Cost.html', {'dept_costs': dept_costs, 'overall_total': overall_total, 'chart_labels': chart_labels, 'chart_values': chart_values})


def pending_payroll_report(request):
    pending_runs = PayrollRun.objects.exclude(status='RELEASED').prefetch_related('lines')
    return render(request, 'Payroll/Pending Payroll.html', {'pending_runs': pending_runs})


def employees_on_leave_report(request):
    from datetime import date
    today = date.today()
    on_leave = LeaveRequest.objects.filter(status='APPROVED', from_date__lte=today, to_date__gte=today).select_related('employee')
    return render(request, 'Payroll/Employees On Leave.html', {'on_leave': on_leave, 'today': today})


def payroll_run_download_docx(request, pk):
    from docx import Document
    run = get_object_or_404(PayrollRun, pk=pk)
    lines = run.lines.select_related('employee').all()
    doc = Document()
    doc.add_heading(f'Payroll Report - {run.month}', level=1)
    doc.add_paragraph(f'Status: {run.get_status_display()}')
    doc.add_paragraph(f'Created: {run.created_at.strftime("%d %b %Y")}')
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Employee Code'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Basic'
    hdr_cells[3].text = 'Gross Salary'
    hdr_cells[4].text = 'Net Pay'
    for line in lines:
        row_cells = table.add_row().cells
        row_cells[0].text = line.employee.employee_code
        row_cells[1].text = line.employee.full_name
        row_cells[2].text = str(line.basic)
        row_cells[3].text = str(line.gross_salary)
        row_cells[4].text = str(line.net_pay)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Payroll_{run.month}.docx"'
    doc.save(response)
    return response


def payroll_run_detail(request, pk):
    run = get_object_or_404(PayrollRun, pk=pk)
    lines = run.lines.select_related('employee').all()
    return render(request, 'Payroll/Payroll Run Detail.html', {'run': run, 'lines': lines})


def payroll_combined(request):
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'create':
            form = PayrollRunForm(request.POST)
            if form.is_valid():
                payroll_run = form.save()
                for emp in Employee.objects.filter(employment_status='ACTIVE'):
                    try:
                        ss = emp.salary_structure
                        PayrollRunLine.objects.create(
                            payroll_run=payroll_run, employee=emp,
                            basic=ss.basic, gross_salary=ss.gross_salary, net_pay=ss.gross_salary,
                        )
                    except SalaryStructure.DoesNotExist:
                        pass
        elif action in ('validate', 'approve', 'release'):
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            next_status = {'validate': 'VALIDATED', 'approve': 'APPROVED', 'release': 'RELEASED'}
            run.status = next_status[action]
            run.save()
        elif action == 'lock':
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            run.is_locked = True
            run.save()
        elif action == 'unlock':
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            run.is_locked = False
            run.save()
        elif action == 'reject':
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            run.status = 'DRAFT'
            run.save()
        elif action == 'reprocess':
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            if run.status == 'RELEASED':
                messages.error(request, f"Cannot reprocess '{run.month}' - it has already been released.")
            elif run.is_locked:
                messages.error(request, f"Cannot reprocess '{run.month}' - it is locked. Unlock it first.")
            else:
                run.lines.all().delete()
                for emp in Employee.objects.filter(employment_status='ACTIVE'):
                    try:
                        ss = emp.salary_structure
                        PayrollRunLine.objects.create(
                            payroll_run=run, employee=emp,
                            basic=ss.basic, gross_salary=ss.gross_salary, net_pay=ss.gross_salary,
                        )
                    except SalaryStructure.DoesNotExist:
                        pass
                messages.success(request, f"'{run.month}' reprocessed successfully.")
        elif action == 'delete':
            run_id = request.POST.get('run_id')
            run = get_object_or_404(PayrollRun, pk=run_id)
            if run.status == 'RELEASED':
                messages.error(request, f"Cannot delete '{run.month}' - it has already been released. Released payroll records must be preserved.")
            else:
                month = run.month
                run.delete()
                messages.success(request, f"'{month}' payroll run deleted.")
        return redirect('payroll_combined')
    else:
        form = PayrollRunForm()
    from django.core.paginator import Paginator
    from django.db.models import Sum, Count
    all_runs = PayrollRun.objects.prefetch_related('lines__employee').all()
    summary = {
        'total_runs': all_runs.count(),
        'total_employees_paid': PayrollRunLine.objects.filter(payroll_run__status='RELEASED').count(),
        'total_net_payout': PayrollRunLine.objects.filter(payroll_run__status='RELEASED').aggregate(t=Sum('net_pay'))['t'] or 0,
        'locked_count': all_runs.filter(is_locked=True).count(),
    }
    paginator = Paginator(all_runs, 10)
    page_number = request.GET.get('page')
    payroll_runs = paginator.get_page(page_number)
    return render(request, 'Payroll/Payroll Combined.html', {'form': form, 'payroll_runs': payroll_runs, 'summary': summary})


def arrears(request):
    from .models import ArrearsRecord
    from .forms import ArrearsForm
    if request.method == 'POST':
        form = ArrearsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('arrears')
    else:
        form = ArrearsForm()
    records = ArrearsRecord.objects.select_related('employee').all()
    return render(request, 'Adjustments/Arrears.html', {'form': form, 'records': records})


def full_final_settlement(request):
    from .models import FullFinalSettlement
    from .forms import FullFinalSettlementForm
    if request.method == 'POST':
        form = FullFinalSettlementForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('full_final_settlement')
    else:
        form = FullFinalSettlementForm()
    records = FullFinalSettlement.objects.select_related('employee').all()
    return render(request, 'Adjustments/Full and Final Settlement.html', {'form': form, 'records': records})


def payslips(request):
    lines = PayrollRunLine.objects.select_related('employee', 'payroll_run').filter(payroll_run__status='RELEASED')
    return render(request, 'Payslips.html', {'lines': lines})


def bank_transfer(request):
    lines = PayrollRunLine.objects.select_related('employee', 'payroll_run').filter(payroll_run__status='RELEASED')
    return render(request, 'Bank Transfer.html', {'lines': lines})


def reports_analytics(request):
    import json
    context = {
        'total_employees': Employee.objects.count(),
        'total_payroll_runs': PayrollRun.objects.count(),
        'total_released': PayrollRun.objects.filter(status='RELEASED').count(),
        'total_pending_leave': LeaveRequest.objects.filter(status='PENDING').count(),
    }
    chart_labels = json.dumps(['Employees', 'Payroll Runs', 'Released', 'Pending Leave'])
    chart_values = json.dumps([context['total_employees'], context['total_payroll_runs'], context['total_released'], context['total_pending_leave']])
    context['chart_labels'] = chart_labels
    context['chart_values'] = chart_values
    return render(request, 'Reports and Analytics.html', context)


def ess(request):
    context = {
        'total_payslips': PayrollRunLine.objects.filter(payroll_run__status='RELEASED').count(),
        'pending_leave': LeaveRequest.objects.filter(status='PENDING').count(),
        'pending_reimbursements': Reimbursement.objects.filter(status='PENDING').count(),
    }
    return render(request, 'ESS.html', context)


def notifications(request):
    from datetime import date, timedelta
    recent_cutoff = date.today() - timedelta(days=30)
    notices = []
    for leave in LeaveRequest.objects.filter(status='PENDING').select_related('employee')[:10]:
        notices.append({'type': 'Leave', 'message': f'{leave.employee.full_name} requested {leave.get_leave_type_display()}', 'date': leave.from_date})
    for reimb in Reimbursement.objects.filter(status='PENDING').select_related('employee')[:10]:
        notices.append({'type': 'Reimbursement', 'message': f'{reimb.employee.full_name} submitted a {reimb.get_category_display()} claim of {reimb.amount}', 'date': reimb.date})
    for run in PayrollRun.objects.filter(status='VALIDATED')[:10]:
        notices.append({'type': 'Payroll', 'message': f'{run.month} payroll is validated and awaiting approval', 'date': run.created_at.date()})
    notices.sort(key=lambda n: n['date'], reverse=True)
    return render(request, 'Notifications.html', {'notices': notices})


def user_roles_permissions(request):
    from .models import UserRoleAssignment
    from .forms import UserRoleAssignmentForm
    if request.method == 'POST':
        form = UserRoleAssignmentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('user_roles_permissions')
    else:
        form = UserRoleAssignmentForm()
    assignments = UserRoleAssignment.objects.select_related('employee').all()
    return render(request, 'User Roles and Permissions.html', {'form': form, 'assignments': assignments})


def settings(request):
    from .models import CompanySettings
    from .forms import CompanySettingsForm
    company_settings, _ = CompanySettings.objects.get_or_create(pk=1)
    if request.method == 'POST':
        form = CompanySettingsForm(request.POST, instance=company_settings)
        if form.is_valid():
            form.save()
            return redirect('settings')
    else:
        form = CompanySettingsForm(instance=company_settings)
    return render(request, 'Settings.html', {'form': form})


def payslip_history(request):
    from django.core.paginator import Paginator
    all_lines = PayrollRunLine.objects.select_related('employee', 'payroll_run').filter(payroll_run__status='RELEASED').order_by('-payroll_run__created_at')
    paginator = Paginator(all_lines, 15)
    page_number = request.GET.get('page')
    lines = paginator.get_page(page_number)
    return render(request, 'payslip management/payslip History.html', {'lines': lines})


def download_pdf(request):
    lines = PayrollRunLine.objects.select_related('employee', 'payroll_run').filter(payroll_run__status='RELEASED')
    return render(request, 'payslip management/Download PDF.html', {'lines': lines})


def _generate_payslip_pdf(line):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from io import BytesIO
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 800, "Payslip")
    c.setFont("Helvetica", 11)
    c.drawString(50, 770, f"Employee: {line.employee.full_name}")
    c.drawString(50, 752, f"Employee Code: {line.employee.employee_code}")
    c.drawString(50, 734, f"Month: {line.payroll_run.month}")
    c.drawString(50, 716, f"Basic: Rs. {line.basic}")
    c.drawString(50, 698, f"Gross Salary: Rs. {line.gross_salary}")
    c.drawString(50, 680, f"Net Pay: Rs. {line.net_pay}")
    c.showPage()
    c.save()
    buf.seek(0)
    return buf.read()


def email_payslip(request):
    from django.core.mail import EmailMessage
    lines = PayrollRunLine.objects.select_related('employee', 'payroll_run').filter(payroll_run__status='RELEASED')
    sent = False
    send_errors = []
    if request.method == 'POST':
        for line in lines:
            if not line.employee.email:
                send_errors.append(f"{line.employee.full_name}: no email on file")
                continue
            try:
                pdf_bytes = _generate_payslip_pdf(line)
                email = EmailMessage(
                    subject=f"Payslip - {line.payroll_run.month}",
                    body=f"Dear {line.employee.full_name},\n\nPlease find attached your payslip for {line.payroll_run.month}.\n\nRegards,\nPayroll Team",
                    to=[line.employee.email],
                )
                email.attach(f"Payslip_{line.employee.employee_code}_{line.payroll_run.month}.pdf", pdf_bytes, "application/pdf")
                email.send(fail_silently=False)
            except Exception as e:
                send_errors.append(f"{line.employee.full_name}: {str(e)}")
        sent = True
        if send_errors:
            messages.error(request, f"{len(send_errors)} email(s) failed: " + "; ".join(send_errors[:5]))
        else:
            messages.success(request, f"Sent {lines.count()} payslip email(s) successfully.")
    return render(request, 'payslip management/Email payslip.html', {'lines': lines, 'sent': sent})


def failed_transaction_report(request):
    return render(request, 'Bank Transfer.html')


def generate_payslip(request):
    return render(request, 'payslip management/Generate payslip.html')


def payment_states(request):
    return render(request, 'Bank Transfer.html')


def salary_transfer_file(request):
    return render(request, 'Bank Transfer.html')


def payslips_export_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="payslips.csv"'
    writer = csv.writer(response)
    writer.writerow(['Month', 'Employee Code', 'Name', 'Basic', 'Gross Salary', 'Net Pay'])
    lines = PayrollRunLine.objects.filter(payroll_run__status='RELEASED').select_related('employee', 'payroll_run')
    for line in lines:
        writer.writerow([
            line.payroll_run.month, line.employee.employee_code, line.employee.full_name,
            line.basic, line.gross_salary, line.net_pay,
        ])
    return response


def payslips_export_excel(request):
    import openpyxl
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Payslips"
    headers = ['Month', 'Employee Code', 'Name', 'Basic', 'Gross Salary', 'Net Pay']
    ws.append(headers)

    lines = PayrollRunLine.objects.filter(payroll_run__status='RELEASED').select_related('employee', 'payroll_run')
    for line in lines:
        ws.append([
            line.payroll_run.month, line.employee.employee_code, line.employee.full_name,
            float(line.basic), float(line.gross_salary), float(line.net_pay),
        ])

    for i, header in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(i)].width = max(14, len(header) + 4)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="payslips.xlsx"'
    wb.save(response)
    return response



def payroll_run_download_excel(request, pk):
    import openpyxl
    from openpyxl.utils import get_column_letter
    run = get_object_or_404(PayrollRun, pk=pk)
    lines = run.lines.select_related('employee').all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Payroll Report'
    ws.append([f'Payroll Report - {run.month}'])
    ws.append([f'Status: {run.get_status_display()}'])
    ws.append([f'Created: {run.created_at.strftime("%d %b %Y")}'])
    ws.append([])
    headers = ['Employee Code', 'Name', 'Basic', 'Gross Salary', 'Net Pay']
    ws.append(headers)
    for line in lines:
        ws.append([
            line.employee.employee_code, line.employee.full_name,
            float(line.basic), float(line.gross_salary), float(line.net_pay),
        ])
    for i, header in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(i)].width = max(16, len(header) + 4)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Payroll_{run.month}.xlsx"'
    wb.save(response)
    return response





from django.contrib.auth import login, logout
from django.contrib.auth.views import LoginView
from django.contrib.auth.decorators import login_required
from .forms import RegisterForm

class CustomLoginView(LoginView):
    template_name = "login.html"

def logout_view(request):
    logout(request)
    return redirect("login")












def request_demo(request):
    from .forms import DemoRequestForm
    if request.method == 'POST':
        form = DemoRequestForm(request.POST)
        if form.is_valid():
            demo_request = form.save()
            try:
                from django.core.mail import send_mail
                from django.conf import settings as django_settings
                send_mail(
                    subject='New Demo Request - ' + demo_request.company_name,
                    message='Full Name: ' + demo_request.full_name + chr(10) + 'Company: ' + demo_request.company_name + chr(10) + 'Email: ' + demo_request.email + chr(10) + 'Phone: ' + demo_request.phone + chr(10) + 'Message: ' + demo_request.message,
                    from_email=django_settings.EMAIL_HOST_USER,
                    recipient_list=['jaganbharath46@gmail.com'],
                    fail_silently=True,
                )
            except Exception:
                pass
            messages.success(request, "Thank you! Your demo request has been received. Our team will contact you shortly.")
            return redirect('request_demo')
    else:
        form = DemoRequestForm()
    return render(request, 'Request Demo.html', {'form': form})


def register_view(request):
    import secrets
    from django.contrib.auth.models import User
    from django.core.mail import send_mail
    from django.conf import settings as django_settings

    error = None
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '')
        password2 = request.POST.get('password2', '')

        if not username:
            error = "Please choose a username."
        elif not email:
            error = "Please enter your email address."
        elif User.objects.filter(username__iexact=username).exists():
            error = "That username is already taken."
        elif User.objects.filter(email__iexact=email).exists():
            error = "That email is already registered."
        elif not EmailOTPVerification.objects.filter(email__iexact=email, is_verified=True).exists():
            error = "Please verify your email address before creating an account."
        elif len(password) < 8:
            error = "Password must be at least 8 characters long."
        elif password != password2:
            error = "Passwords do not match."
        else:
            user = User.objects.create_user(username=username, email=email, password=password)
            user.is_active = False
            user.save()
            try:
                send_mail(
                    subject='Registration received - Payroll Management',
                    message=f'Hi {username},\n\nYour email has been verified and your registration was received. An administrator will review and approve your account before you can log in.',
                    from_email=django_settings.EMAIL_HOST_USER,
                    recipient_list=[user.email],
                    fail_silently=True,
                )
                messages.success(request, "Registration submitted! Please check your email to verify your account. After verification, an admin will approve your access.")
            except Exception as e:
                messages.warning(request, f"Account created but verification email failed to send: {str(e)}. Contact admin.")
            return redirect('login')

    return render(request, 'register.html', {'error': error})


def verify_email(request, token):
    try:
        record = EmailVerificationToken.objects.get(token=token, is_used=False)
    except EmailVerificationToken.DoesNotExist:
        messages.error(request, "Invalid or expired verification link.")
        return redirect('login')

    record.is_used = True
    record.email_verified = True
    record.save()
    messages.success(request, "Email verified successfully! Your account is now pending admin approval. You'll be able to log in once approved.")
    return redirect('login')


def landing_page(request):
    return render(request, 'Landing.html')










def send_verification_code(request):
    import json
    import random
    from django.contrib.auth.hashers import make_password
    from django.contrib.auth.models import User
    from django.core.mail import send_mail
    from django.conf import settings as django_settings
    from django.utils import timezone
    from datetime import timedelta
    from django.http import JsonResponse

    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Invalid request method.'}, status=405)

    try:
        data = json.loads(request.body)
    except Exception:
        data = request.POST

    email = data.get('email', '').strip().lower()

    if not email or '@' not in email or '.' not in email.split('@')[-1]:
        return JsonResponse({'success': False, 'message': 'Please enter a valid email address.'})

    if User.objects.filter(email__iexact=email).exists():
        return JsonResponse({'success': False, 'message': 'This email is already registered. Please sign in instead.'})

    existing = EmailOTPVerification.objects.filter(email__iexact=email).order_by('-last_sent_at').first()
    if existing and (timezone.now() - existing.last_sent_at).total_seconds() < 60:
        wait = 60 - int((timezone.now() - existing.last_sent_at).total_seconds())
        return JsonResponse({'success': False, 'message': f'Please wait {wait} seconds before requesting another code.'})

    code = str(random.randint(100000, 999999))
    otp_record = EmailOTPVerification.objects.create(
        email=email,
        otp_hash=make_password(code),
        expires_at=timezone.now() + timedelta(minutes=10),
        attempts=0,
        is_verified=False,
    )

    try:
        send_mail(
            subject='Your verification code - Payroll Management',
            message=f'Your verification code is: {code}\n\nThis code will expire in 10 minutes.\n\nIf you did not request this, please ignore this email.',
            from_email=django_settings.EMAIL_HOST_USER,
            recipient_list=[email],
            fail_silently=False,
        )
        return JsonResponse({'success': True, 'message': 'Verification code sent to your email.'})
    except Exception as e:
        otp_record.delete()
        return JsonResponse({'success': False, 'message': 'Failed to send verification email. Please try again.'})




def verify_code(request):
    import json
    from django.contrib.auth.hashers import check_password
    from django.utils import timezone
    from django.http import JsonResponse

    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Invalid request method.'}, status=405)

    try:
        data = json.loads(request.body)
    except Exception:
        data = request.POST

    email = data.get('email', '').strip().lower()
    code = data.get('code', '').strip()

    if not email or not code:
        return JsonResponse({'success': False, 'message': 'Email and code are required.'})

    otp_record = EmailOTPVerification.objects.filter(email__iexact=email, is_verified=False).order_by('-created_at').first()

    if not otp_record:
        return JsonResponse({'success': False, 'message': 'No pending verification found for this email. Please request a new code.'})

    if timezone.now() > otp_record.expires_at:
        return JsonResponse({'success': False, 'message': 'This code has expired. Please request a new one.'})

    if otp_record.attempts >= 5:
        return JsonResponse({'success': False, 'message': 'Too many incorrect attempts. Please request a new code.'})

    if not check_password(code, otp_record.otp_hash):
        otp_record.attempts += 1
        otp_record.save()
        remaining = 5 - otp_record.attempts
        return JsonResponse({'success': False, 'message': f'Incorrect code. {remaining} attempt(s) remaining.'})

    otp_record.is_verified = True
    otp_record.save()
    return JsonResponse({'success': True, 'message': 'Email verified successfully.'})
